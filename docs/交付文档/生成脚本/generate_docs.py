#!/usr/bin/env python3
"""生成 StudyMate 的 DOCX 与 PDF 部署交付文档。

封面和视觉样式沿用参考文档，正文按逻辑块自然分页；目录页码由
最终正文分页动态生成，不再通过固定 44/9 页制造留白或整页缩字。
"""

from __future__ import annotations

import html
import math
import textwrap
from pathlib import Path

import qrcode
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

from doc_content import (
    AUTHORS,
    DOMAIN,
    ICP,
    LONG_PAGES,
    PROJECT_TITLE,
    SHORT_PAGES,
)


SCRIPT_DIR = Path(__file__).resolve().parent
DELIVERY_DIR = SCRIPT_DIR.parent
OUTPUT_DIR = DELIVERY_DIR / "交付文件"
ASSETS = SCRIPT_DIR / "assets"

BODY_FONT_FILE = Path(
    "/home/ysc/.local/share/fonts/Microsoft/OfficeCompatibility/仿宋_GB2312.ttf"
)
HEADING_FONT_FILE = Path(
    "/home/ysc/.local/share/fonts/Microsoft/OfficeCompatibility/simhei.ttf"
)
CODE_FONT_FILE = Path(
    "/home/ysc/.local/share/fonts/Developer/MapleMono/MapleMono-NF-CN-Regular.ttf"
)

DOC_BODY_FONT = "仿宋_GB2312"
DOC_HEADING_FONT = "黑体"
DOC_CODE_FONT = "Maple Mono NF CN"

BLUE = "365F91"
LIGHT_BLUE = "DCE6F1"
VERY_LIGHT_BLUE = "F4F8FC"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"

CONTINUATION_SUFFIX = "（续）"


def is_continuation_heading(block) -> bool:
    return (
        block.get("kind") == "heading"
        and str(block.get("text", "")).endswith(CONTINUATION_SUFFIX)
    )


def compatible_adjacent_tables(left, right) -> bool:
    """Merge tables that were split only to satisfy the former fixed pages."""
    if left.get("kind") != "table" or right.get("kind") != "table":
        return False
    if left.get("merge_columns") or right.get("merge_columns"):
        return False
    keys = ("headers", "widths", "font_size", "center_columns", "width_ratio")
    return all(left.get(key) == right.get(key) for key in keys)


def logical_blocks(pages) -> list[dict]:
    """Flatten legacy page groups into a continuous logical document stream."""
    blocks: list[dict] = []
    for page in pages:
        for source in page:
            if is_continuation_heading(source):
                continue
            block = dict(source)
            if blocks and compatible_adjacent_tables(blocks[-1], block):
                blocks[-1]["rows"] = [*blocks[-1]["rows"], *block["rows"]]
                continue
            blocks.append(block)
    return blocks


def toc_entries(blocks) -> list[tuple[int, str]]:
    """Return level-1/2 headings in logical order for DOCX and PDF TOCs."""
    return [
        (block["level"], block["text"])
        for block in blocks
        if block.get("kind") == "heading" and block.get("level") in (1, 2)
    ]


def bookmark_map(entries) -> dict[str, tuple[str, int]]:
    return {
        title: (f"toc_{index:04d}", index + 1)
        for index, (_level, title) in enumerate(entries)
    }


def load_pil_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = "#17365D",
    line_gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=font)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + line_gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    outline: str = "#4F81BD",
    text_fill: str = "#17365D",
    radius: int = 22,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=4)
    draw_centered_text(draw, box, text, font, text_fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#4F81BD",
    width: int = 7,
) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 20
    for delta in (2.55, -2.55):
        p = (
            end[0] + size * math.cos(angle + delta),
            end[1] + size * math.sin(angle + delta),
        )
        draw.line([end, p], fill=color, width=width)


def generate_architecture_long(path: Path) -> None:
    image = PILImage.new("RGB", (1800, 930), "white")
    draw = ImageDraw.Draw(image)
    title = load_pil_font(HEADING_FONT_FILE, 42)
    main = load_pil_font(HEADING_FONT_FILE, 34)
    small = load_pil_font(BODY_FONT_FILE, 28)

    draw.text((60, 36), "StudyMate 生产部署架构", font=title, fill="#17365D")
    rounded_box(draw, (620, 120, 1180, 230), "用户浏览器\nHTTPS / SSE", main, "#EAF2F8")
    rounded_box(draw, (620, 300, 1180, 410), "Caddy 公网网关\n80 / 443 · 自动 HTTPS", main, "#D9EAF7")
    rounded_box(draw, (620, 480, 1180, 590), "Frontend Nginx\nReact 静态站点 · /api 代理", main, "#E2F0D9", "#70AD47")
    rounded_box(draw, (620, 660, 1180, 780), "FastAPI Backend\n认证 · 智能体 · RAG · SSE", main, "#FFF2CC", "#BF9000")

    rounded_box(draw, (80, 690, 480, 850), "SQLite 数据库\nbackend_data\n种子 34 账号 · 5 课程\n938 知识块 / 向量", small, "#FCE4D6", "#C55A11")
    rounded_box(draw, (1320, 690, 1720, 850), "Piston 沙箱\nPython / C / C++\n2 CPU · 2 GB · 2 并发", small, "#E4DFEC", "#8064A2")
    rounded_box(draw, (1320, 300, 1720, 450), "外部服务\n大模型 · 讯飞语音\n公开学习资源", small, "#F2F2F2", "#7F7F7F")

    arrow(draw, (900, 230), (900, 300))
    arrow(draw, (900, 410), (900, 480))
    arrow(draw, (900, 590), (900, 660))
    arrow(draw, (620, 750), (480, 770))
    arrow(draw, (1180, 750), (1320, 770))
    arrow(draw, (1180, 700), (1410, 450))

    draw.text((1230, 140), "matropic.cn → 121.40.64.199", font=small, fill="#595959")
    draw.text((85, 885), "公网仅开放 22 / 80 / 443；5173 / 8000 / 2000 仅绑定 127.0.0.1", font=small, fill="#595959")
    image.save(path, quality=95)


def generate_architecture_short(path: Path) -> None:
    image = PILImage.new("RGB", (1200, 1500), "white")
    draw = ImageDraw.Draw(image)
    title = load_pil_font(HEADING_FONT_FILE, 46)
    main = load_pil_font(HEADING_FONT_FILE, 36)
    small = load_pil_font(BODY_FONT_FILE, 29)

    draw.text((70, 45), "StudyMate 系统架构", font=title, fill="#17365D")
    rounded_box(draw, (150, 145, 1050, 275), "访问层：浏览器 / PC / 移动端", main, "#EAF2F8")
    rounded_box(draw, (150, 365, 1050, 510), "公网网关：Caddy\nDNS · HTTPS · HTTP/2 · 压缩", main, "#D9EAF7")
    rounded_box(draw, (150, 600, 1050, 745), "展示层：React + Nginx\n静态资源 · SPA · /api 反向代理", main, "#E2F0D9", "#70AD47")
    rounded_box(draw, (150, 835, 1050, 995), "应用层：FastAPI + Agent Orchestrator\n认证 · 课程 · RAG · SSE · 多智能体", main, "#FFF2CC", "#BF9000")
    rounded_box(draw, (80, 1110, 535, 1300), "数据层\nSQLite · BM25\nbackend_data", small, "#FCE4D6", "#C55A11")
    rounded_box(draw, (665, 1110, 1120, 1300), "能力层\nPiston · LLM API\nPython / C / C++", small, "#E4DFEC", "#8064A2")

    arrow(draw, (600, 275), (600, 365))
    arrow(draw, (600, 510), (600, 600))
    arrow(draw, (600, 745), (600, 835))
    arrow(draw, (500, 995), (310, 1110))
    arrow(draw, (700, 995), (890, 1110))
    draw.text((90, 1400), "Docker Compose 统一编排 · Named Volume 持久化 · UFW 与安全组双重防护", font=small, fill="#595959")
    image.save(path, quality=95)


def generate_deployment_flow(path: Path) -> None:
    image = PILImage.new("RGB", (1800, 620), "white")
    draw = ImageDraw.Draw(image)
    title = load_pil_font(HEADING_FONT_FILE, 40)
    font = load_pil_font(HEADING_FONT_FILE, 28)
    small = load_pil_font(BODY_FONT_FILE, 23)
    draw.text((55, 35), "StudyMate 部署流程", font=title, fill="#17365D")
    labels = [
        ("1 服务器与 DNS", "Ubuntu 22.04\nA 记录 / 安全组"),
        ("2 安装 Docker", "国内 CE 源\n镜像加速 / UFW"),
        ("3 上传与配置", "rsync\n.deploy.env / .env"),
        ("4 构建与启动", "Compose\nPiston runtime"),
        ("5 HTTPS 验收", "Caddy 证书\n登录 / AI / 代码"),
    ]
    box_w, gap, y1, y2 = 285, 62, 170, 480
    x = 50
    for idx, (heading, body) in enumerate(labels):
        draw.rounded_rectangle((x, y1, x + box_w, y2), radius=24, fill="#EEF4FB", outline="#4F81BD", width=4)
        draw_centered_text(draw, (x + 15, y1 + 25, x + box_w - 15, y1 + 105), heading, font)
        draw.line((x + 30, y1 + 125, x + box_w - 30, y1 + 125), fill="#9FBAD0", width=3)
        draw_centered_text(draw, (x + 15, y1 + 140, x + box_w - 15, y2 - 20), body, small, "#404040")
        if idx < len(labels) - 1:
            arrow(draw, (x + box_w + 8, (y1 + y2) // 2), (x + box_w + gap - 8, (y1 + y2) // 2))
        x += box_w + gap
    image.save(path, quality=95)


def generate_mobile_web(path: Path) -> None:
    image = PILImage.new("RGB", (900, 1180), "white")
    draw = ImageDraw.Draw(image)
    title = load_pil_font(HEADING_FONT_FILE, 40)
    main = load_pil_font(HEADING_FONT_FILE, 29)
    small = load_pil_font(BODY_FONT_FILE, 24)

    draw.text((55, 35), "StudyMate 响应式移动 Web", font=title, fill="#333333")
    draw.rounded_rectangle((185, 120, 715, 1080), radius=54, fill="#202020", outline="#202020", width=4)
    draw.rounded_rectangle((215, 170, 685, 1015), radius=30, fill="#FAFAFA", outline="#BFBFBF", width=3)
    draw.rounded_rectangle((330, 137, 570, 165), radius=14, fill="#555555")
    draw.rounded_rectangle((245, 205, 655, 265), radius=20, fill="#F2F2F2", outline="#BFBFBF", width=2)
    draw.text((275, 220), "https://matropic.cn", font=small, fill="#333333")
    draw.text((255, 310), "StudyMate", font=main, fill="#17365D")
    draw.text((255, 355), "随时继续你的学习闭环", font=small, fill="#595959")
    cards = [
        ("五门课程", "课程隔离 · RAG 来源追溯"),
        ("AI 助教", "SSE 流式回答 · 附件 · 语音"),
        ("学习工作台", "笔记 · 测验 · 报告 · 路径"),
        ("账号同步", "与 PC 共用 HTTPS 与后端数据"),
    ]
    y = 425
    for heading, detail in cards:
        draw.rounded_rectangle((250, y, 650, y + 120), radius=18, fill="#FFFFFF", outline="#A6A6A6", width=2)
        draw.text((275, y + 18), heading, font=main, fill="#333333")
        draw.text((275, y + 68), detail, font=small, fill="#666666")
        y += 140
    draw.text((235, 1040), "无需安装 · 无独立小程序 · PC 端功能更完整", font=small, fill="#666666")
    image.save(path, quality=95)


def generate_assets() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    generate_architecture_long(ASSETS / "architecture_long.png")
    generate_architecture_short(ASSETS / "architecture_short.png")
    generate_deployment_flow(ASSETS / "deployment_flow.png")
    generate_mobile_web(ASSETS / "mobile_web.png")
    qr = qrcode.QRCode(version=5, box_size=16, border=3)
    qr.add_data(DOMAIN)
    qr.make(fit=True)
    qr.make_image(fill_color="#17365D", back_color="white").convert("RGB").save(ASSETS / "site_qr.png")


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="A6A6A6", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    set_table_borders(table, color="FFFFFF", size="0")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, east_asia: str, size: float, bold=False, color=None) -> None:
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if east_asia != DOC_CODE_FONT else DOC_CODE_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman" if east_asia != DOC_CODE_FONT else DOC_CODE_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_compact(paragraph, before=0, after=0, line=1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, DOC_BODY_FONT, 9)


def request_field_update(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_pageref_field(paragraph, bookmark: str, fallback_page: int, size: float) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = str(fallback_page)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, DOC_BODY_FONT, size)


def restart_page_number(section, start=1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def unlink_headers_and_footers(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def set_section_size(section, long_doc: bool) -> None:
    if long_doc:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Cm(1.75)
        section.bottom_margin = Cm(1.55)
        section.left_margin = Cm(2.15)
        section.right_margin = Cm(2.05)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.05)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.65)


def configure_doc_styles(doc: Document, long_doc: bool) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = DOC_BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_BODY_FONT)
    normal.font.size = Pt(10 if long_doc else 10.2)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(2)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = DOC_HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_HEADING_FONT)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].font.size = Pt(16 if long_doc else 17)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(0)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 2"].font.size = Pt(13 if long_doc else 13.5)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(4)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(3)
    doc.styles["Heading 3"].font.size = Pt(11 if long_doc else 11.5)
    doc.styles["Heading 3"].font.bold = True
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(3)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(2)


def add_cover(doc: Document, long_doc: bool) -> None:
    if long_doc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(ASSETS / "university_seal.jpg"), width=Cm(1.6))
        p.add_run("   ").add_picture(str(ASSETS / "university_name.jpg"), width=Cm(4.9))

        for _ in range(3):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("StudyMate：基于大模型的")
        set_run_font(run, DOC_HEADING_FONT, 18, True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("个性化资源生成与学习多智能体系统")
        set_run_font(run, DOC_HEADING_FONT, 18, True)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("项目部署文档")
        set_run_font(run, DOC_HEADING_FONT, 30, True)

        for _ in range(5):
            doc.add_paragraph()
        info = doc.add_table(rows=3, cols=2)
        info.alignment = 1
        remove_table_borders(info)
        labels = ["作品名称：", "申报者姓名：", "集体申报全体成员姓名："]
        values = [PROJECT_TITLE, "________", AUTHORS]
        for row, label, value in zip(info.rows, labels, values):
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            r1 = row.cells[0].paragraphs[0].add_run(label)
            r2 = row.cells[1].paragraphs[0].add_run(value)
            set_run_font(r1, DOC_HEADING_FONT, 12, True)
            set_run_font(r2, DOC_BODY_FONT, 10 if label == "作品名称：" else 12)
            set_paragraph_compact(row.cells[0].paragraphs[0], after=8)
            set_paragraph_compact(row.cells[1].paragraphs[0], after=8)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(PROJECT_TITLE)
        set_run_font(run, DOC_HEADING_FONT, 15, True, GRAY)
        p.paragraph_format.space_after = Pt(0)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:color"), "7F7F7F")
        borders.append(bottom)
        p_pr.append(borders)

        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"＜{PROJECT_TITLE}＞")
        set_run_font(run, DOC_BODY_FONT, 12.5, False, GRAY)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("系统部署说明")
        set_run_font(run, DOC_HEADING_FONT, 31, True)

        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"作者：{AUTHORS}")
        set_run_font(run, DOC_BODY_FONT, 13)


def add_toc_line(
    doc: Document,
    level: int,
    title: str,
    page: int,
    bookmark: str,
    compact: bool,
) -> None:
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0.65 if level == 2 else 0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0 if compact else 1)
    fmt.line_spacing = 1.0
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9200")
    tabs.append(tab)
    paragraph._p.get_or_add_pPr().append(tabs)
    run = paragraph.add_run(title)
    font_size = 8.1 if compact else 10
    set_run_font(run, DOC_HEADING_FONT if level == 1 else DOC_BODY_FONT, font_size, level == 1)
    tab_run = paragraph.add_run("\t")
    set_run_font(tab_run, DOC_BODY_FONT, font_size)
    add_pageref_field(paragraph, bookmark, page, font_size)


def add_toc(doc: Document, toc, pages, bookmarks, long_doc: bool) -> None:
    if not long_doc:
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(PROJECT_TITLE)
        set_run_font(run, DOC_HEADING_FONT, 8.5, True, GRAY)
        p_pr = header._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "7F7F7F")
        borders.append(bottom)
        p_pr.append(borders)
        header.paragraph_format.space_after = Pt(14)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("目  录")
    set_run_font(run, DOC_HEADING_FONT, 20 if long_doc else 22, True)
    title.paragraph_format.space_after = Pt(8 if long_doc else 16)
    for level, text_value in toc:
        bookmark, _bookmark_id = bookmarks[text_value]
        add_toc_line(
            doc,
            level,
            text_value,
            pages.get(text_value, 1),
            bookmark,
            compact=long_doc,
        )


def add_doc_heading(doc: Document, block, long_doc: bool, bookmarks=None) -> None:
    level = min(3, block["level"])
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(block["text"])
    set_run_font(
        run,
        DOC_HEADING_FONT,
        {1: 16 if long_doc else 17, 2: 13 if long_doc else 13.5, 3: 11 if long_doc else 11.5}[level],
        True,
    )
    if bookmarks and block["text"] in bookmarks:
        name, bookmark_id = bookmarks[block["text"]]
        add_bookmark(paragraph, name, bookmark_id)


def add_doc_paragraph(doc: Document, block, long_doc: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.7)
    paragraph.paragraph_format.widow_control = True
    set_paragraph_compact(paragraph, after=2, line=1.08 if long_doc else 1.1)
    text_value = block["text"]
    prefix = block.get("bold_prefix")
    if prefix and text_value.startswith(prefix):
        r = paragraph.add_run(prefix)
        set_run_font(r, DOC_HEADING_FONT, 9.7 if long_doc else 10, True)
        r = paragraph.add_run(text_value[len(prefix):])
        set_run_font(r, DOC_BODY_FONT, 9.7 if long_doc else 10)
    else:
        r = paragraph.add_run(text_value)
        set_run_font(r, DOC_BODY_FONT, 9.7 if long_doc else 10)


def add_doc_bullets(doc: Document, block, long_doc: bool) -> None:
    for idx, item in enumerate(block["items"], 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        set_paragraph_compact(paragraph, after=1, line=1.02)
        prefix = f"{idx}. " if block.get("ordered") else "• "
        r = paragraph.add_run(prefix + item)
        set_run_font(r, DOC_BODY_FONT, 9.3 if long_doc else 9.7)


def add_doc_code(doc: Document, block, long_doc: bool) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    set_table_borders(tbl, color="BFBFBF", size="4")
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=55, start=85, bottom=55, end=85)
    paragraph = cell.paragraphs[0]
    set_paragraph_compact(paragraph, after=0, line=1.0)
    for idx, line in enumerate(block["text"].splitlines()):
        if idx:
            paragraph.add_run().add_break()
        r = paragraph.add_run(line or " ")
        set_run_font(r, DOC_CODE_FONT, 6.9 if long_doc else 7.2)
    spacer_p = doc.add_paragraph()
    set_paragraph_compact(spacer_p, after=0)
    spacer_p.paragraph_format.line_spacing = Pt(2)


def table_merge_ranges(block):
    """Return contiguous same-value ranges for requested data columns."""
    ranges = []
    data_rows = block["rows"]
    for column in block.get("merge_columns", []):
        start = 0
        while start < len(data_rows):
            value = str(data_rows[start][column])
            end = start
            while end + 1 < len(data_rows) and str(data_rows[end + 1][column]) == value:
                end += 1
            if value and end > start:
                ranges.append((column, start, end, value))
            start = end + 1
    return ranges


def add_doc_table(doc: Document, block, long_doc: bool) -> None:
    rows = [block["headers"]] + block["rows"]
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.alignment = 1
    tbl.autofit = False
    set_table_borders(tbl, color="A6A6A6", size="4")
    if block.get("widths"):
        for row in tbl.rows:
            for cell, width in zip(row.cells, block["widths"]):
                cell.width = Cm(width)
    center_columns = set(block.get("center_columns", []))
    for ridx, row_values in enumerate(rows):
        row = tbl.rows[ridx]
        if ridx == 0:
            set_repeat_table_header(row)
        for cidx, value in enumerate(row_values):
            cell = row.cells[cidx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=45, start=55, bottom=45, end=55)
            if ridx == 0:
                set_cell_shading(cell, "E7E6E6")
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if ridx == 0 or cidx in center_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_compact(p, after=0, line=1.0)
            r = p.add_run(str(value))
            set_run_font(r, DOC_HEADING_FONT if ridx == 0 else DOC_BODY_FONT, block.get("font_size", 8.5), ridx == 0)
    for column, start, end, value in table_merge_ranges(block):
        merged = tbl.cell(start + 1, column).merge(tbl.cell(end + 1, column))
        merged.text = ""
        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(merged, "F2F2F2")
        set_cell_margins(merged, top=45, start=55, bottom=45, end=55)
        paragraph = merged.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_compact(paragraph, after=0, line=1.0)
        run = paragraph.add_run(value)
        set_run_font(run, DOC_BODY_FONT, block.get("font_size", 8.5))
    spacer_p = doc.add_paragraph()
    set_paragraph_compact(spacer_p, after=0)
    spacer_p.paragraph_format.line_spacing = Pt(2)


def add_doc_image(doc: Document, block) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = ASSETS / block["asset"]
    kwargs = {"width": Cm(block["width_cm"])}
    if block.get("height_cm"):
        kwargs["height"] = Cm(block["height_cm"])
    paragraph.add_run().add_picture(str(path), **kwargs)
    set_paragraph_compact(paragraph, after=0)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(caption, after=3)
    r = caption.add_run(block["caption"])
    set_run_font(r, DOC_BODY_FONT, 8.5)


def add_doc_note(doc: Document, block, long_doc: bool) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    set_table_borders(tbl, color="BFBFBF", size="5")
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=65, start=90, bottom=65, end=90)
    p = cell.paragraphs[0]
    set_paragraph_compact(p, after=0, line=1.05)
    r = p.add_run(f"{block['title']}：")
    set_run_font(r, DOC_HEADING_FONT, 9 if long_doc else 9.5, True)
    r = p.add_run(block["text"])
    set_run_font(r, DOC_BODY_FONT, 9 if long_doc else 9.5)
    spacer_p = doc.add_paragraph()
    set_paragraph_compact(spacer_p, after=0)
    spacer_p.paragraph_format.line_spacing = Pt(2)


def render_doc_block(doc: Document, block, long_doc: bool, bookmarks=None) -> None:
    kind = block["kind"]
    if kind == "heading":
        add_doc_heading(doc, block, long_doc, bookmarks)
    elif kind == "paragraph":
        add_doc_paragraph(doc, block, long_doc)
    elif kind == "bullets":
        add_doc_bullets(doc, block, long_doc)
    elif kind == "code":
        add_doc_code(doc, block, long_doc)
    elif kind == "table":
        add_doc_table(doc, block, long_doc)
    elif kind == "image":
        add_doc_image(doc, block)
    elif kind == "note":
        add_doc_note(doc, block, long_doc)
    elif kind == "spacer":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(block["points"])
    else:
        raise ValueError(f"Unsupported block kind: {kind}")


def configure_body_header_footer(section, long_doc: bool) -> None:
    unlink_headers_and_footers(section)
    section.header.paragraphs[0].clear()
    if not long_doc:
        p = section.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(PROJECT_TITLE)
        set_run_font(r, DOC_HEADING_FONT, 8.5, True, GRAY)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "7F7F7F")
        borders.append(bottom)
        p_pr.append(borders)
    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_field(footer.paragraphs[0])


def make_docx(path: Path, long_doc: bool, heading_pages=None) -> None:
    pages = LONG_PAGES if long_doc else SHORT_PAGES
    blocks = logical_blocks(pages)
    toc = toc_entries(blocks)
    bookmarks = bookmark_map(toc)
    heading_pages = heading_pages or {}
    doc = Document()
    doc.core_properties.title = "项目部署文档" if long_doc else "系统部署说明书"
    doc.core_properties.subject = PROJECT_TITLE
    doc.core_properties.author = "StudyMate 项目组"
    doc.core_properties.comments = "按竞赛示例版式生成；作者信息请由项目组补充。"
    set_section_size(doc.sections[0], long_doc)
    configure_doc_styles(doc, long_doc)
    request_field_update(doc)

    add_cover(doc, long_doc)
    doc.add_page_break()
    add_toc(doc, toc, heading_pages, bookmarks, long_doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_size(body_section, long_doc)
    configure_body_header_footer(body_section, long_doc)
    restart_page_number(body_section, 1)

    for block in blocks:
        render_doc_block(doc, block, long_doc, bookmarks)

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------


def register_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont("FangSong", str(BODY_FONT_FILE)))
    pdfmetrics.registerFont(TTFont("SimHei", str(HEADING_FONT_FILE)))
    pdfmetrics.registerFont(TTFont("MapleMono", str(CODE_FONT_FILE)))


def pdf_styles(long_doc: bool) -> dict[str, ParagraphStyle]:
    body_size = 10.2 if long_doc else 10.0
    return {
        "h1": ParagraphStyle(
            "h1", fontName="SimHei", fontSize=16.5 if long_doc else 16.5,
            leading=20, spaceAfter=7, textColor=colors.black, wordWrap="CJK",
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "h2", fontName="SimHei", fontSize=13 if long_doc else 13,
            leading=16, spaceBefore=4, spaceAfter=4, wordWrap="CJK",
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "h3", fontName="SimHei", fontSize=11 if long_doc else 11,
            leading=14, spaceBefore=3, spaceAfter=3, wordWrap="CJK",
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "body", fontName="FangSong", fontSize=body_size,
            leading=body_size * (1.36 if long_doc else 1.28), firstLineIndent=body_size * 2,
            alignment=TA_JUSTIFY, spaceAfter=4 if long_doc else 3, wordWrap="CJK"
        ),
        "bullet": ParagraphStyle(
            "bullet", fontName="FangSong", fontSize=body_size - 0.2,
            leading=body_size * (1.28 if long_doc else 1.2), leftIndent=14, firstLineIndent=-10,
            spaceAfter=2 if long_doc else 1, wordWrap="CJK"
        ),
        "caption": ParagraphStyle(
            "caption", fontName="FangSong", fontSize=8.8 if long_doc else 8.2,
            leading=11, alignment=TA_CENTER, spaceAfter=4, wordWrap="CJK"
        ),
        "note": ParagraphStyle(
            "note", fontName="FangSong", fontSize=9.4 if long_doc else 9.1,
            leading=12, wordWrap="CJK"
        ),
    }


def escaped(text_value) -> str:
    return html.escape(str(text_value)).replace("\n", "<br/>")


def wrap_code(text_value: str, width: int = 95) -> str:
    result = []
    for line in text_value.splitlines() or [""]:
        if len(line) <= width:
            result.append(line)
            continue
        indent = len(line) - len(line.lstrip())
        chunks = textwrap.wrap(
            line,
            width=width,
            subsequent_indent=" " * min(indent + 2, 12),
            break_long_words=True,
            break_on_hyphens=False,
            replace_whitespace=False,
        )
        result.extend(chunks or [line])
    return "\n".join(result)


def pdf_block(block, styles, long_doc: bool, avail_width: float, bookmarks=None):
    kind = block["kind"]
    if kind == "heading":
        paragraph = Paragraph(
            escaped(block["text"]),
            styles[f"h{min(3, block['level'])}"],
        )
        if bookmarks and block["text"] in bookmarks:
            bookmark, _bookmark_id = bookmarks[block["text"]]
            paragraph._toc_level = block["level"] - 1
            paragraph._toc_text = block["text"]
            paragraph._bookmark = bookmark
        return [paragraph]
    if kind == "paragraph":
        text_value = escaped(block["text"])
        prefix = block.get("bold_prefix")
        if prefix and block["text"].startswith(prefix):
            text_value = f"<font name='SimHei'>{escaped(prefix)}</font>{escaped(block['text'][len(prefix):])}"
        return [Paragraph(text_value, styles["body"])]
    if kind == "bullets":
        flow = []
        for idx, item in enumerate(block["items"], 1):
            marker = f"{idx}." if block.get("ordered") else "•"
            flow.append(Paragraph(f"{marker}&nbsp;&nbsp;{escaped(item)}", styles["bullet"]))
        return flow
    if kind == "code":
        code_style = ParagraphStyle(
            "code", fontName="MapleMono", fontSize=7.4 if long_doc else 7,
            leading=9.2 if long_doc else 8.7, textColor=colors.HexColor("#202020")
        )
        # Preformatted treats its input as literal text and preserves real line
        # breaks, so HTML escaping would leak entities such as &lt; into output.
        pre = Preformatted(wrap_code(block["text"], width=112 if long_doc else 108), code_style)
        tbl = Table([[pre]], colWidths=[avail_width - 4])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3F3F3")),
            ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#BFBFBF")),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return [tbl, Spacer(1, 3)]
    if kind == "table":
        fs = block.get("font_size", 8.5) * (1.12 if long_doc else 1.0)
        cell_style = ParagraphStyle("cell", fontName="FangSong", fontSize=fs, leading=fs * 1.15, wordWrap="CJK")
        center_cell_style = ParagraphStyle(
            "cellCenter", parent=cell_style, alignment=TA_CENTER
        )
        head_style = ParagraphStyle("head", fontName="SimHei", fontSize=fs, leading=fs * 1.15, alignment=TA_CENTER, wordWrap="CJK")
        data = [[Paragraph(escaped(v), head_style) for v in block["headers"]]]
        center_columns = set(block.get("center_columns", []))
        data.extend([
            [
                Paragraph(escaped(v), center_cell_style if cidx in center_columns else cell_style)
                for cidx, v in enumerate(row)
            ]
            for row in block["rows"]
        ])
        col_widths = None
        if block.get("widths"):
            total = sum(block["widths"])
            table_width = avail_width * block.get("width_ratio", 1.0)
            col_widths = [table_width * w / total for w in block["widths"]]
        tbl = Table(data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
        table_style = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E7E6E6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("BOX", (0, 0), (-1, -1), 0.65, colors.HexColor("#737373")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#A6A6A6")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4 if long_doc else 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4 if long_doc else 3),
        ]
        for column, start, end, _value in table_merge_ranges(block):
            table_style.append(("SPAN", (column, start + 1), (column, end + 1)))
            table_style.append(("ALIGN", (column, start + 1), (column, end + 1), "CENTER"))
            table_style.append(("BACKGROUND", (column, start + 1), (column, end + 1), colors.HexColor("#F2F2F2")))
        tbl.setStyle(TableStyle(table_style))
        return [tbl, Spacer(1, 3)]
    if kind == "image":
        path = ASSETS / block["asset"]
        width = block["width_cm"] * cm
        if block.get("height_cm"):
            height = block["height_cm"] * cm
        else:
            with PILImage.open(path) as image:
                height = width * image.height / image.width
        image_flow = RLImage(str(path), width=width, height=height)
        image_flow.hAlign = "CENTER"
        caption = Paragraph(escaped(block["caption"]), styles["caption"])
        return [KeepTogether([image_flow, caption])]
    if kind == "note":
        content = Paragraph(
            f"<font name='SimHei'>{escaped(block['title'])}：</font>{escaped(block['text'])}",
            styles["note"],
        )
        tbl = Table([[content]], colWidths=[avail_width - 4])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F2F2F2")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFBFBF")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return [tbl, Spacer(1, 3)]
    if kind == "spacer":
        return [Spacer(1, block["points"])]
    raise ValueError(f"Unsupported block kind: {kind}")


def draw_pdf_cover(c, page_size, long_doc: bool) -> None:
    w, h = page_size
    if long_doc:
        seal = ASSETS / "university_seal.jpg"
        school = ASSETS / "university_name.jpg"
        c.drawImage(str(seal), 55, h - 80, width=43, height=41, preserveAspectRatio=True, mask="auto")
        c.drawImage(str(school), 108, h - 69, width=118, height=24, preserveAspectRatio=True, mask="auto")
        c.setFont("SimHei", 18)
        c.setFillColor(colors.black)
        c.drawCentredString(w / 2, h - 195, "StudyMate：基于大模型的")
        c.drawCentredString(w / 2, h - 230, "个性化资源生成与学习多智能体系统")
        c.setFont("SimHei", 29)
        c.drawCentredString(w / 2, h - 315, "项目部署文档")

        label_x = 118
        value_x = 230
        y = 270
        for label, value in (
            ("作品名称：", PROJECT_TITLE),
            ("申报者姓名：", "________"),
            ("集体申报全体成员姓名：", AUTHORS),
        ):
            c.setFont("SimHei", 11)
            c.drawRightString(label_x + 100, y, label)
            value_size = 9.5 if label == "作品名称：" else 11
            c.setFont("FangSong", value_size)
            if len(value) > 30:
                p = Paragraph(escaped(value), ParagraphStyle("coverValue", fontName="FangSong", fontSize=value_size, leading=12, wordWrap="CJK"))
                p.wrapOn(c, w - value_x - 45, 40)
                p.drawOn(c, value_x, y - 4)
            else:
                c.drawString(value_x, y, value)
            y -= 44
    else:
        c.setStrokeColor(colors.HexColor("#7F7F7F"))
        c.setLineWidth(1.5)
        c.line(70, h - 75, w - 70, h - 75)
        c.setFillColor(colors.HexColor("#666666"))
        c.setFont("SimHei", 14)
        c.drawCentredString(w / 2, h - 57, PROJECT_TITLE)
        c.setFillColor(colors.HexColor("#666666"))
        cover_subtitle = Paragraph(
            escaped(f"＜{PROJECT_TITLE}＞"),
            ParagraphStyle(
                "shortCoverProject",
                fontName="FangSong",
                fontSize=12.5,
                leading=16,
                alignment=TA_CENTER,
                wordWrap="CJK",
            ),
        )
        cover_subtitle.wrapOn(c, w - 150, 45)
        cover_subtitle.drawOn(c, 75, h - 225)
        c.setFillColor(colors.black)
        c.setFont("SimHei", 31)
        c.drawCentredString(w / 2, h - 270, "系统部署说明")
        c.setFillColor(colors.black)
        c.setFont("FangSong", 12)
        c.drawCentredString(w / 2, h - 335, f"作者：{AUTHORS}")


def draw_body_header_footer(c, page_size, number: int, long_doc: bool) -> None:
    w, h = page_size
    c.setFillColor(colors.HexColor("#555555"))
    c.setFont("FangSong", 8.5)
    c.drawCentredString(w / 2, 22, str(number))
    if not long_doc:
        c.setFillColor(colors.HexColor("#666666"))
        c.setFont("SimHei", 8.3)
        c.drawCentredString(w / 2, h - 29, PROJECT_TITLE)
        c.setStrokeColor(colors.HexColor("#7F7F7F"))
        c.setLineWidth(0.7)
        c.line(58, h - 37, w - 58, h - 37)


def draw_pdf_toc_page_header(c, page_size, long_doc: bool) -> None:
    if long_doc:
        return
    w, h = page_size
    c.setFillColor(colors.HexColor("#666666"))
    c.setFont("SimHei", 8.3)
    c.drawCentredString(w / 2, h - 28, PROJECT_TITLE)
    c.setStrokeColor(colors.HexColor("#7F7F7F"))
    c.setLineWidth(0.7)
    c.line(58, h - 37, w - 58, h - 37)


class StudyMatePDFDocument(BaseDocTemplate):
    def __init__(self, path: Path, long_doc: bool, styles) -> None:
        self.long_doc = long_doc
        self.styles_map = styles
        self.page_size = LETTER if long_doc else A4
        self.body_page = 0
        self.heading_pages: dict[str, int] = {}
        super().__init__(
            str(path),
            pagesize=self.page_size,
            pageCompression=1,
            title="项目部署文档" if long_doc else "系统部署说明书",
            author="StudyMate 项目组",
            subject=PROJECT_TITLE,
            leftMargin=0,
            rightMargin=0,
            topMargin=0,
            bottomMargin=0,
        )

        w, h = self.page_size
        cover_frame = Frame(0, 0, w, h, id="cover-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
        if long_doc:
            toc_frame = Frame(66, 48, w - 132, h - 92, id="toc-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
            body_frame = Frame(58, 38, w - 116, h - 82, id="body-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
        else:
            toc_frame = Frame(66, 82, w - 132, h - 138, id="toc-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
            body_frame = Frame(60, 38, w - 120, h - 91, id="body-frame", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

        def cover_page(canv, _doc):
            draw_pdf_cover(canv, self.page_size, self.long_doc)

        def toc_page(canv, _doc):
            draw_pdf_toc_page_header(canv, self.page_size, self.long_doc)

        def body_page(canv, _doc):
            self.body_page += 1
            draw_body_header_footer(canv, self.page_size, self.body_page, self.long_doc)

        self.addPageTemplates([
            PageTemplate(id="cover", frames=[cover_frame], onPage=cover_page),
            PageTemplate(id="toc", frames=[toc_frame], onPage=toc_page),
            PageTemplate(id="body", frames=[body_frame], onPage=body_page),
        ])

    def beforeDocument(self) -> None:
        self.body_page = 0
        self.heading_pages = {}

    def afterFlowable(self, flowable) -> None:
        if not hasattr(flowable, "_toc_level"):
            return
        level = flowable._toc_level
        text_value = flowable._toc_text
        bookmark = flowable._bookmark
        self.canv.bookmarkPage(bookmark)
        self.canv.addOutlineEntry(text_value, bookmark, level=level, closed=False)
        self.notify("TOCEntry", (level, text_value, self.body_page, bookmark))
        self.heading_pages[text_value] = self.body_page


def build_pdf_story(pages, long_doc: bool, styles):
    blocks = logical_blocks(pages)
    entries = toc_entries(blocks)
    bookmarks = bookmark_map(entries)

    toc = TableOfContents()
    toc.dotsMinLevel = 0
    toc.levelStyles = [
        ParagraphStyle(
            "toc-level-1",
            fontName="SimHei",
            fontSize=7.8 if long_doc else 10.2,
            leading=9.4 if long_doc else 13,
            leftIndent=0,
            firstLineIndent=0,
            spaceAfter=0,
        ),
        ParagraphStyle(
            "toc-level-2",
            fontName="FangSong",
            fontSize=7.4 if long_doc else 9.8,
            leading=9.2 if long_doc else 12.5,
            leftIndent=15,
            firstLineIndent=0,
            spaceAfter=0,
        ),
    ]
    toc_title = ParagraphStyle(
        "toc-title",
        fontName="SimHei",
        fontSize=20 if long_doc else 22,
        leading=25,
        alignment=TA_CENTER,
        spaceAfter=9 if long_doc else 16,
    )

    story = [NextPageTemplate("toc"), PageBreak()]
    story.extend([
        Paragraph("目&nbsp;&nbsp;录", toc_title),
        toc,
        NextPageTemplate("body"),
        PageBreak(),
    ])
    for block in blocks:
        story.extend(pdf_block(block, styles, long_doc, (LETTER if long_doc else A4)[0] - (116 if long_doc else 120), bookmarks))
    return story


def make_pdf(path: Path, long_doc: bool):
    pages = LONG_PAGES if long_doc else SHORT_PAGES
    styles = pdf_styles(long_doc)
    doc = StudyMatePDFDocument(path, long_doc, styles)
    doc.multiBuild(build_pdf_story(pages, long_doc, styles))
    return dict(doc.heading_pages), doc.body_page


def main() -> None:
    for font_path in (BODY_FONT_FILE, HEADING_FONT_FILE, CODE_FONT_FILE):
        if not font_path.exists():
            raise FileNotFoundError(font_path)
    for source_asset in (ASSETS / "university_seal.jpg", ASSETS / "university_name.jpg"):
        if not source_asset.exists():
            raise FileNotFoundError(source_asset)

    register_pdf_fonts()
    generate_assets()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    outputs = [
        ("项目部署文档", True),
        ("系统部署说明书", False),
    ]
    for stem, long_doc in outputs:
        pdf_path = OUTPUT_DIR / f"{stem}.pdf"
        docx_path = OUTPUT_DIR / f"{stem}.docx"
        print(f"Generating {pdf_path.name} ...", flush=True)
        heading_pages, body_pages = make_pdf(pdf_path, long_doc)
        print(f"  natural body pages: {body_pages}", flush=True)
        print(f"Generating {docx_path.name} ...", flush=True)
        make_docx(docx_path, long_doc, heading_pages)
    print("Done.")


if __name__ == "__main__":
    main()
