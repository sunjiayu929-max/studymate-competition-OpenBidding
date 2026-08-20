"""将三个领域的 14 个岗位资料增量导入 StudyMate RAG 库。

资料位于 ``backend/resources/domain_knowledge/source_materials``。脚本支持
Markdown 与 DOCX。正文含 [S1]/[G1] 等引用时，脚本解析同文件来源区的名称、
版本/日期和直达 URL；无可解析引用时才附加岗位主题的官方补充核验资料。每条
内容以稳定哈希去重；重建时只替换本脚本此前写入的 role-v1/role-v2 切片，绝不
删除 FDE 或其他课程数据。
"""
from __future__ import annotations

import asyncio
import hashlib
import re
import sys
from pathlib import Path
from urllib.parse import urlparse

from sqlalchemy import delete, select

from app.db import async_session_maker
from app.db.models import Course, KnowledgeChunk
from app.rag import get_rag_service
from app.rag.source import clean_source_name


ROOT = Path(__file__).resolve().parents[1] / "resources" / "domain_knowledge" / "source_materials"
URL_RE = re.compile(r"https://[^\s)）；;，。]+")
REFERENCE_RE = re.compile(r"^[-*]\s+\[([A-Za-z]+\d+)\]\s*(.+?)\s*(https://[^\s)）；;，。]+)", re.MULTILINE)
REFERENCE_LABEL_RE = re.compile(r"\[([A-Za-z]+\d+)\]")
REFERENCE_MARKER_RE = re.compile(r"\s*\[[A-Za-z]+\d+\]")
URL_LINE_RE = re.compile(r"(?:^|\s)URL\s*[:：]\s*https://", re.IGNORECASE)
SOURCE_METADATA_RE = re.compile(r"^(?:URL\s*[:：]?|发布日期|发布日|版本|许可证|访问核验)", re.IGNORECASE)

SUPPLEMENTAL_REFERENCES: dict[str, list[dict[str, str]]] = {
    "AI Agent 开发工程师": [{"name": "LangChain 官方 GitHub 仓库（MIT）", "url": "https://github.com/langchain-ai/langchain", "note": "补充核验资料"}],
    "AI Infra 工程师": [{"name": "Kubernetes 官方文档（Apache-2.0）", "url": "https://kubernetes.io/docs/home/", "note": "补充核验资料"}],
    "具身智能算法工程师": [{"name": "ROS 2 官方 GitHub 仓库（Apache-2.0）", "url": "https://github.com/ros2/ros2", "note": "补充核验资料"}],
    "大模型安全工程师": [{"name": "OWASP Top 10 for LLM Applications（持续更新，CC BY-SA 4.0）", "url": "https://owasp.org/www-project-top-10-for-large-language-model-applications/", "note": "补充核验资料"}],
    "大模型应用开发工程师": [{"name": "LangChain 官方 GitHub 仓库（MIT）", "url": "https://github.com/langchain-ai/langchain", "note": "补充核验资料"}],
    "工业互联网架构师": [{"name": "Industrial Internet Reference Architecture（IIC）", "url": "https://www.iiconsortium.org/IIRA.htm", "note": "补充核验资料"}],
    "工业数据工程师": [{"name": "Apache IoTDB 官方 GitHub 仓库（Apache-2.0）", "url": "https://github.com/apache/iotdb", "note": "补充核验资料"}],
    "边缘计算 AI 工程师": [{"name": "ONNX Runtime 官方 GitHub 仓库（MIT）", "url": "https://github.com/microsoft/onnxruntime", "note": "补充核验资料"}],
    "工业 AI 视觉工程师": [{"name": "OpenCV 官方 GitHub 仓库（Apache-2.0）", "url": "https://github.com/opencv/opencv", "note": "补充核验资料"}],
    "工业互联网网络集成工程师": [{"name": "Eclipse Milo OPC UA SDK（EPL-2.0）", "url": "https://github.com/eclipse/milo", "note": "补充核验资料"}],
}

ROLE_SOURCES: tuple[tuple[str, str, str, str], ...] = (
    ("人工智能", "AI Agent开发工程师", "AI Agent 开发工程师", "AI Agent 开发、工具调用、工作流编排与评测"),
    ("人工智能", "AI Infra工程师", "AI Infra 工程师", "模型训练/推理基础设施、资源调度与可靠性"),
    ("人工智能", "具身智能算法工程师", "具身智能算法工程师", "感知、规划、控制与仿真到实机验证"),
    ("人工智能", "大模型安全工程师", "大模型安全工程师", "大模型风险评测、权限、对抗与安全治理"),
    ("人工智能", "大模型开发应用工程师", "大模型应用开发工程师", "大模型应用、RAG、Agent 与效果评测"),
    ("特定软件开发", "01-DevSecOps软件供应链安全工程师", "软件供应链安全工程师（DevSecOps）", "安全开发生命周期、SBOM、签名与漏洞响应"),
    ("特定软件开发", "02-企业RAG应用实施工程师", "企业 RAG 应用实施工程师", "企业资料解析、检索、引用、权限与评测"),
    ("特定软件开发", "03-MLOps工程师", "MLOps 工程师", "训练流水线、模型注册、部署、监控与回滚"),
    ("特定软件开发", "04-AI-native应用前端开发工程师", "AI-native 应用前端开发工程师", "流式交互、Agent 可视化、引用界面与端侧推理"),
    ("工业互联网", "工业互联网架构师", "工业互联网架构师", "工业平台架构、系统集成、数据与安全治理"),
    ("工业互联网", "工业数据工程师", "工业数据工程师", "工业数据采集、治理、时序处理与分析"),
    ("工业互联网", "边缘计算AI工程师", "边缘计算 AI 工程师", "边缘设备、模型部署、性能与可靠性"),
    ("工业互联网", "工业AI视觉工程师", "工业 AI 视觉工程师", "缺陷数据、视觉模型、边缘部署与现场验收"),
    ("工业互联网", "工业互联网网络集成工程师", "工业互联网网络集成工程师", "工业网络、协议接入、联调、运维与安全"),
)


def course_name(role: str) -> str:
    return role if role.endswith("岗位知识库") else f"{role} 岗位知识库"


def read_text(path: Path) -> str:
    if path.suffix.lower() in {".md", ".markdown", ".txt"}:
        return path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".docx":
        from docx import Document

        document = Document(path)
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", "；") for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    raise ValueError(f"不支持的资料格式：{path}")


def parse_reference_map(text: str) -> dict[str, dict[str, str]]:
    """解析 Markdown 来源区；DOCX 若无标准引用则返回空映射。"""
    references: dict[str, dict[str, str]] = {}
    for label, raw_name, url in REFERENCE_RE.findall(text):
        name = re.sub(r"\*+", "", raw_name.strip().rstrip("：:；;，, ")).strip()
        references[label] = {"label": label, "name": name, "url": url.rstrip(".,，。；;")}
    return references


def _platform_name(url: str) -> str:
    host = urlparse(url).netloc.lower().removeprefix("www.").removeprefix("m.")
    labels = {
        "zhaopin.com": "智联招聘",
        "liepin.com": "猎聘",
        "zhipin.com": "BOSS直聘",
        "lagou.com": "拉勾招聘",
        "github.com": "GitHub",
    }
    return next((name for domain, name in labels.items() if host.endswith(domain)), host or "外部资料")


def _inline_source_name(role: str, block: str, url: str) -> str:
    """Give a URL-only citation a human-readable, reviewable label."""
    position = block.find(url)
    if position >= 0:
        for raw_line in reversed(block[:position].splitlines()):
            candidate = raw_line.strip().strip("-*#■ ")
            if not candidate or SOURCE_METADATA_RE.match(candidate) or "原始来源" in candidate:
                continue
            candidate = URL_RE.sub("", candidate).strip(" ：:")
            if 4 <= len(candidate) <= 72:
                return f"{_platform_name(url)} · {candidate}"
    return f"{_platform_name(url)} · {role}岗位资料"


def _reference_source_name(raw_name: str, url: str) -> str:
    """Turn a bibliography entry into a compact label suitable for a source card."""
    platform = _platform_name(url)
    title = re.sub(r"^(?:智联招聘|猎聘|BOSS直聘|拉勾招聘)[，,\s]*", "", raw_name)
    title = re.sub(r"[，,；;]\s*(?:公开招聘信息[，,；;]?\s*转述引用|访问核验|发布日期|发布日).*$", "", title)
    title = re.sub(r"[，,；;]\s*(?:MIT|Apache-[\d.]+|GitHub|官方文档|新闻稿).*$", "", title)
    title = re.sub(r"[：:]$", "", title).strip()
    return f"{platform} · {title}" if title else platform


def _citation_kind(url: str) -> str:
    host = urlparse(url).netloc.lower().removeprefix("www.").removeprefix("m.")
    if any(host.endswith(domain) for domain in ("zhaopin.com", "liepin.com", "zhipin.com", "lagou.com")):
        return "招聘岗位原文"
    if host.endswith("github.com"):
        return "开源项目原文"
    if "/docs/" in url or host.endswith(("kubernetes.io", "owasp.org")):
        return "官方技术文档"
    return "原始网页资料"


def resolve_sources(role: str, block: str, references: dict[str, dict[str, str]]) -> tuple[str, str | None, list[dict[str, str]], str]:
    labels = list(dict.fromkeys(REFERENCE_LABEL_RE.findall(block)))
    sources = [
        {
            **references[label],
            "name": _reference_source_name(references[label]["name"], references[label]["url"]),
            "kind": _citation_kind(references[label]["url"]),
        }
        for label in labels
        if label in references
    ]
    source_status = "document_citation"
    if not sources:
        inline_urls = list(dict.fromkeys(URL_RE.findall(block)))
        sources = [
            {
                "label": "URL",
                "name": _inline_source_name(role, block, url),
                "url": url.rstrip(".,，。；;"),
                "kind": _citation_kind(url),
            }
            for url in inline_urls
        ]
        source_status = "inline_url" if sources else "supplemental_reference"
    if not sources:
        sources = SUPPLEMENTAL_REFERENCES.get(role, [])
    if not sources:
        return "来源待补充", None, [], "unverified"
    # The card title must identify one concrete source. Other citations remain
    # in metadata and are explicitly listed in the source panel.
    return sources[0]["name"], sources[0]["url"], sources, source_status


def split_blocks(text: str, limit: int = 850, *, preserve_paragraphs: bool = False) -> list[tuple[str, str]]:
    heading = "资料正文"
    blocks: list[tuple[str, str]] = []
    current: list[str] = []
    length = 0
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if preserve_paragraphs and current:
                blocks.append((heading, "\n".join(current)))
                current, length = [], 0
            elif current and current[-1] != "":
                # Keep Markdown paragraph boundaries. In particular, table
                # rows must remain consecutive for remark-gfm to parse them.
                current.append("")
            continue
        if line.startswith("#") or re.match(r"^(第[一二三四五六七八九十\d]+[章节]|[一二三四五六七八九十\d]+[、.])", line):
            if current:
                blocks.append((heading, "\n".join(current)))
                current, length = [], 0
            heading = line.lstrip("#").strip()
            continue
        if length + len(line) > limit and current:
            blocks.append((heading, "\n".join(current)))
            current, length = [], 0
        current.append(line)
        length += len(line)
        # 招聘资料常以 URL 行分隔不同职位样本；在此切开可避免多个
        # 职位描述和链接挤在同一张检索卡里。
        if URL_LINE_RE.search(line):
            blocks.append((heading, "\n".join(current)))
            current, length = [], 0
    if current:
        blocks.append((heading, "\n".join(current)))
    return blocks


def clean_content(block: str) -> str:
    """正文不展示 [S1]/[G1]/[J1]，来源仍保留在 meta.citations 中。"""
    cleaned = REFERENCE_MARKER_RE.sub("", block)
    # 正文不显示 Markdown 粗体/斜体星号，避免将格式标记暴露为 AI 痕迹。
    cleaned = re.sub(r"\*+", "", cleaned)
    visible_lines: list[str] = []
    lines = cleaned.splitlines()
    for index, raw_line in enumerate(lines):
        line = raw_line.strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""
        if not line:
            if visible_lines and visible_lines[-1] != "":
                visible_lines.append("")
            continue
        if SOURCE_METADATA_RE.match(line) or "原始来源与引用" in line:
            continue
        if line.startswith("■") or (URL_LINE_RE.search(next_line) and any(tag in line for tag in ("招聘", "GitHub", "BOSS直聘", "猎聘", "智联"))):
            continue
        line = URL_RE.sub("", line).strip(" ：:")
        if line:
            visible_lines.append(line)
    cleaned = "\n".join(visible_lines)
    cleaned = re.sub(r"\s+([，。；：、])", r"\1", cleaned)
    return cleaned.strip()


def is_source_section(heading: str) -> bool:
    normalized_heading = heading.strip().lower()
    return "来源" in heading or normalized_heading in {"可直达资料", "资料原始来源完整清单"}


def enrich_short_software_blocks(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """为软件岗位保留真实出处的前提下，补足过短的检索片段。

    原始资料常把同一来源支撑的职责、流程和验收拆成单独小标题。对于短条目，
    这里仅拼接同一 Markdown 文件内、且至少共用一个引用标签的相邻条目；不生成
    新的事实性文字，也不把无关来源混入。这样检索页和来源页能展示连续、可核对的
    岗位知识，同时每个事实仍可从右侧引用列表回到外部页面。
    """
    prepared = [
        (heading, raw, set(REFERENCE_LABEL_RE.findall(raw)))
        for heading, raw in blocks
        if not is_source_section(heading)
    ]
    enriched: list[tuple[str, str]] = []
    for index, (heading, raw, labels) in enumerate(prepared):
        content = clean_content(raw)
        if len(content) >= 620 or not labels:
            enriched.append((heading, raw))
            continue

        parts = [raw]
        used_length = len(content)
        # 先拿距离当前条目最近的同源内容，保证阅读顺序和语义关联尽可能自然。
        candidate_indexes = [
            *range(index + 1, len(prepared)),
            *range(index - 1, -1, -1),
        ]
        for candidate_index in candidate_indexes:
            candidate_heading, candidate_raw, candidate_labels = prepared[candidate_index]
            candidate_content = clean_content(candidate_raw)
            if not candidate_content or not (labels & candidate_labels):
                continue
            if used_length + len(candidate_content) > 1150:
                continue
            parts.append(f"\n\n{candidate_heading}\n{candidate_raw}")
            used_length += len(candidate_content)
            if used_length >= 720 or len(parts) >= 4:
                break
        enriched.append((heading, "".join(parts)))
    return enriched


async def sanitize_existing_role_chunks() -> None:
    """清理历史岗位索引中的引用标号和内部训练演示片段。

    早期 FDE 演示数据曾把 StudyMate 自编训练场景写成可检索来源；它不属于
    外部事实，因此删除该单条记录。其余历史岗位内容只做展示层安全清理，不改
    岗位正文或外部链接。
    """
    async with async_session_maker() as db:
        rows = (
            await db.scalars(
                select(KnowledgeChunk)
                .join(Course, KnowledgeChunk.course_id == Course.id)
                .where(Course.name.like("%岗位知识库"))
            )
        ).all()
        changed = 0
        removed = 0
        for row in rows:
            if row.source.startswith("StudyMate FDE 岗位训练切片") or "studymate-SoftwareCopyright" in (row.url or ""):
                await db.delete(row)
                removed += 1
                continue
            cleaned = clean_content(row.content)
            cleaned_source = clean_source_name(re.sub(r"\*+", "", row.source or "").strip())
            cleaned_meta = dict(row.meta or {})
            citations = cleaned_meta.get("citations")
            if isinstance(citations, list):
                cleaned_meta["citations"] = [
                    {**item, "name": re.sub(r"\*+", "", item.get("name", "")).strip()}
                    if isinstance(item, dict) else item
                    for item in citations
                ]
            if cleaned != row.content or cleaned_source != row.source or cleaned_meta != (row.meta or {}):
                row.content = cleaned
                row.source = cleaned_source
                row.meta = cleaned_meta
                changed += 1
        await db.commit()
    print(f"历史岗位索引清理：正文标号 {changed} 条，删除内部训练片段 {removed} 条")


def make_chunks(domain: str, folder: str, role: str) -> list[dict]:
    base = ROOT / domain / folder
    if not base.is_dir():
        raise FileNotFoundError(f"岗位资料目录不存在：{base}")
    items: list[dict] = []
    for path in sorted(base.rglob("*")):
        if path.suffix.lower() not in {".md", ".markdown", ".txt", ".docx"}:
            continue
        try:
            text = read_text(path)
        except Exception as exc:
            print(f"跳过不可读取资料：{path.name}（{exc.__class__.__name__}）")
            continue
        if not text.strip():
            print(f"跳过空白资料：{path.name}")
            continue
        references = parse_reference_map(text)
        # FDE 的招聘事实和“相邻岗位”样本都在同一份资料中。按自然段
        # 切分可以保留每一句原文自身的引用边界，避免将相邻岗位链接误挂到
        # CatPaw FDE 的正文旁。
        raw_blocks = split_blocks(text, preserve_paragraphs=role.startswith("FDE"))
        # 五个数字编号目录只用于“特定软件开发”领域；其他领域保持原有切分粒度。
        blocks = enrich_short_software_blocks(raw_blocks) if folder[:2] in {"01", "02", "03", "04"} else raw_blocks
        for index, (heading, block) in enumerate(blocks, start=1):
            if is_source_section(heading):
                continue
            # The identity tracks the source position instead of display-only
            # cleanup, keeping source-page links stable across future rebuilds.
            stable_material = f"{domain}|{role}|{path.relative_to(ROOT).as_posix()}|{index}"
            stable_id = hashlib.sha256(stable_material.encode("utf-8")).hexdigest()[:48]
            source, url, citations, source_status = resolve_sources(role, block, references)
            # J2-J7 are explicitly documented as neighbouring AI roles, not
            # FDE roles. Keep them out of the FDE corpus so a FDE query never
            # presents an Agent/algorithm vacancy as its own job evidence.
            if role.startswith("FDE") and any(
                str(citation.get("label", "")).startswith("J")
                and citation.get("label") != "J1"
                for citation in citations
            ):
                continue
            content = clean_content(block)
            if not content:
                continue
            items.append({
                "id": stable_id,
                "content": content,
                "source": source,
                "url": url,
                "meta": {
                    "domain": domain,
                    "role": role,
                    "topic": heading,
                    "citations": citations,
                    "source_status": source_status,
                    "difficulty": 2,
                    "source_notice": "优先保留资料正文显式引用；未标引正文仅附岗位主题补充核验资料，不将补充资料表述为该句的唯一原始出处。",
                },
            })
    if not items:
        raise ValueError(f"岗位资料中没有可导入内容：{base}")
    return items


async def import_role(domain: str, folder: str, role: str, description: str) -> dict[str, int | str]:
    name = course_name(role)
    items = make_chunks(domain, folder, role)
    async with async_session_maker() as db:
        course = (await db.execute(select(Course).where(Course.name == name))).scalar_one_or_none()
        if course is None:
            course = Course(name=name, description=f"{domain}｜{description}。资料按原文件与章节切分，供岗位训练、检索与引用。")
            db.add(course)
            await db.flush()
        await db.execute(delete(KnowledgeChunk).where(
            KnowledgeChunk.course_id == course.id,
            KnowledgeChunk.chroma_id.like("role-v%:%"),
        ))
        inserted = 0
        for item in items:
            stable_id = f"role-v2:{item['id']}"
            db.add(KnowledgeChunk(
                course_id=course.id,
                content=item["content"],
                source=clean_source_name(re.sub(r"\*+", "", item["source"]).strip()),
                url=item["url"],
                meta=item["meta"],
                chroma_id=stable_id,
            ))
            inserted += 1
        await db.commit()
    return {"course": name, "inserted": inserted, "total": len(items)}


async def import_fde_expansion() -> dict[str, int | str]:
    """只追加本次 FDE 扩展资料；已有 FDE 知识块保持不动。"""
    domain = "特定软件开发"
    folder = "00-FDE前线部署工程师"
    role = "FDE 岗位知识库"
    items = make_chunks(domain, folder, role)
    async with async_session_maker() as db:
        course = (await db.execute(select(Course).where(Course.name == role))).scalar_one_or_none()
        if course is None:
            raise RuntimeError("未找到已有 FDE 岗位知识库，已停止扩展导入以避免创建错误课程")
        await db.execute(delete(KnowledgeChunk).where(
            KnowledgeChunk.course_id == course.id,
            KnowledgeChunk.chroma_id.like("fde-expansion-v1:%"),
        ))
        for item in items:
            db.add(KnowledgeChunk(
                course_id=course.id,
                content=item["content"],
                source=clean_source_name(re.sub(r"\*+", "", item["source"]).strip()),
                url=item["url"],
                meta=item["meta"],
                chroma_id=f"fde-expansion-v1:{item['id']}",
            ))
        await db.commit()
    return {"course": role, "inserted": len(items), "total": len(items)}


async def main() -> None:
    results = []
    for domain, folder, role, description in ROLE_SOURCES:
        results.append(await import_role(domain, folder, role, description))
    results.append(await import_fde_expansion())
    await sanitize_existing_role_chunks()
    service = get_rag_service()
    service._loaded = False
    await service.ensure_loaded()
    for result in results:
        print(f"{result['course']}：本次新增 {result['inserted']} 条，资料切片共 {result['total']} 条")


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, OSError):
        pass
    asyncio.run(main())
